from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "software-quality" / "availability-testing-learnings.docx"

BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREEN = "E2F0D9"
LIGHT_YELLOW = "FFF2CC"
BORDER = "4F4F4F"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, margin=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin))
        node.set(qn("w:type"), "dxa")


def style_table(table, header_fill=LIGHT_BLUE):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if row_index == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 78, 121)


def set_widths(table, widths):
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)


def add_table(doc, headers, rows, widths=None, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            row.cells[index].text = value
    if widths:
        set_widths(table, widths)
    style_table(table, header_fill)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(31, 78, 121)
    return paragraph


def add_callout(doc, title, text, fill):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    set_cell_border(cell, "A6A6A6")
    set_cell_margins(cell, 180)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    paragraph.add_run(f"\n{text}")
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Availability Testing: Changes And Lessons Learned")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Software Quality - Availability, replication, failover, and manual validation")
    sub_run.italic = True
    sub_run.font.size = Pt(11)
    doc.add_paragraph()

    add_callout(
        doc,
        "Purpose",
        "This document summarizes the practical changes and lessons learned from the availability unit tests and the manual tests performed against the running application. The focus is on how testing shaped the design, not only on whether the final behavior passed.",
        LIGHT_GREEN,
    )

    add_heading(doc, "1. Scope", 1)
    doc.add_paragraph(
        "The work covers the availability part of the system: SQL primary database, SQL secondary database, application-level replication queue, retry behavior, failover, failback, and write blocking during transition states."
    )
    doc.add_paragraph(
        "The current implementation is deliberately scoped for a Software Quality project. It demonstrates availability and recovery behavior without trying to implement enterprise-grade high availability or distributed transactions."
    )

    add_heading(doc, "2. Key Changes Driven By Testing", 1)
    add_table(
        doc,
        ["Area", "Change", "Why it was changed"],
        [
            (
                "Datasource setup",
                "Primary and secondary datasource beans are now configured explicitly. The primary datasource and primary JdbcTemplate are marked as primary/default.",
                "Manual startup testing showed that adding a secondary datasource could confuse Spring Boot/JPA unless the primary datasource was made explicit.",
            ),
            (
                "Replication trigger",
                "Replication jobs are processed by an immediate internal event and by a scheduled worker as fallback.",
                "Manual testing showed that waiting only for a scheduled worker made replication feel too passive. The event gives a fast first attempt while the scheduler preserves retry behavior.",
            ),
            (
                "Retry timing",
                "Failed jobs now get a nextAttemptAt delay before they are retried.",
                "Manual testing showed that retries could happen too quickly and move jobs to dead-letter almost immediately when secondary was down.",
            ),
            (
                "Primary write success check",
                "Account create/update now uses saveAndFlush before creating a replication job.",
                "Manual duplicate-key testing showed that replication must only start after primary persistence has truly succeeded.",
            ),
            (
                "Queue robustness",
                "The in-memory queue guards against adding the same job id to completed or dead-letter more than once.",
                "During manual observation, completed counts were checked closely. The guard makes the observable state more robust if concurrent triggers overlap.",
            ),
        ],
        [1.35, 3.1, 3.15],
    )

    add_heading(doc, "3. Automated Unit Test Feedback", 1)
    doc.add_paragraph(
        "The availability unit tests are written as focused tests using Arrange, Act, Assert. Mocks are used for side effects such as health probes, replication gateways, and replication service calls."
    )
    add_table(
        doc,
        ["Test focus", "Representative test", "What it helped verify"],
        [
            (
                "State routing",
                "DatabaseRoutingServiceTest",
                "The state machine starts in PRIMARY_ACTIVE, switches to SECONDARY_ACTIVE during failover, supports failback, and rejects invalid transitions.",
            ),
            (
                "Health threshold",
                "PrimaryHealthServiceTest",
                "Primary is only considered unavailable after the configured number of consecutive failures, and a healthy check resets the failure count.",
            ),
            (
                "Write coordination",
                "WriteOperationCoordinatorTest",
                "Writes to primary create replication jobs, writes on secondary do not replicate back to secondary, failed primary writes do not enqueue jobs, and transition states block writes.",
            ),
            (
                "Replication worker",
                "ReplicationWorkerTest",
                "Successful jobs complete, failing jobs are retried, and jobs move to dead-letter after the retry limit.",
            ),
            (
                "Account secondary replication",
                "JdbcAccountSecondaryReplicationGatewayTest",
                "Account CREATE, UPDATE, and DELETE operations map to the expected SQL for the secondary database.",
            ),
            (
                "Failover/failback services",
                "FailoverAndFailbackServiceTest",
                "Manual failover requires the queue to drain, emergency failover attempts to drain, failback requires a healthy primary, and sync is called before returning to primary.",
            ),
        ],
        [1.5, 2.0, 4.1],
    )

    add_heading(doc, "4. Manual Test Results", 1)
    add_table(
        doc,
        ["Manual test", "Expected behavior", "Observed result / learning"],
        [
            (
                "Create account while primary is active",
                "Primary write succeeds, replication job is created, and secondary is updated.",
                "Confirmed. The account replication flow works for the phase 1 account scope.",
            ),
            (
                "Delete account while primary is active",
                "Primary delete succeeds and a DELETE replication job removes the account from secondary.",
                "Confirmed. A create followed by delete correctly counts as two completed replication jobs.",
            ),
            (
                "Duplicate account create",
                "Primary rejects the request and no replication job should be created.",
                "Testing revealed the need for saveAndFlush so database constraint failure happens before replication is queued.",
            ),
            (
                "Secondary unavailable",
                "Primary request still succeeds, replication fails, and the job becomes observable as pending/retry/dead-letter.",
                "Confirmed. Testing also revealed that retry attempts needed a delay to avoid moving to dead-letter too quickly.",
            ),
            (
                "Manual failover with pending jobs",
                "Failover is rejected with 503 because secondary is missing at least one primary write.",
                "Confirmed. The response showed: Manual failover requires an empty replication queue.",
            ),
            (
                "Delete while failback/maintenance is active",
                "Writes and deletes are blocked while the system is in a transition state.",
                "Confirmed. The expected service-unavailable error was returned.",
            ),
        ],
        [1.75, 2.65, 3.2],
    )

    add_heading(doc, "5. Important Lessons Learned", 1)
    add_table(
        doc,
        ["Lesson", "Explanation"],
        [
            (
                "Observable state matters",
                "The status endpoint made manual testing much easier because queue counts, active role, maintenance mode, and health failure count could be inspected directly.",
            ),
            (
                "Asynchronous replication needs careful timing",
                "Immediate processing improves responsiveness, but retry behavior needs delay and limits to avoid burning through retries too quickly.",
            ),
            (
                "Primary success must mean persisted success",
                "A returned object from JPA is not always enough for replication decisions. saveAndFlush makes database constraint failures visible before the replication job is created.",
            ),
            (
                "State transitions are not just documentation",
                "Manual testing confirmed that FAILOVER_IN_PROGRESS and FAILBACK_IN_PROGRESS affect behavior by blocking writes and deletes.",
            ),
            (
                "Manual tests found integration issues unit tests would not catch alone",
                "Startup datasource configuration, CORS/runtime behavior, and real secondary database failures were discovered through running the application.",
            ),
        ],
        [2.0, 5.6],
        LIGHT_YELLOW,
    )

    add_heading(doc, "6. Current Limitations", 1)
    doc.add_paragraph(
        "The solution is intentionally limited. The queue is in-memory, so jobs are lost if the application stops. Only account replication is implemented in phase 1. The normal JPA repository flow still uses the primary datasource; full runtime repository routing to secondary is not implemented yet. Failback synchronization is represented as a service boundary and is suitable for demonstration, but it is not a complete production data reconciliation system."
    )

    add_heading(doc, "7. Suggested Next Tests", 1)
    add_table(
        doc,
        ["Next test level", "Suggested test"],
        [
            ("Endpoint tests", "Verify /availability/status, /availability/failover, /availability/failback/begin, and /availability/failback/complete responses."),
            ("Integration tests", "Run with a real secondary SQL database and verify account create/delete replication end to end."),
            ("Failure integration test", "Stop secondary, create an account, verify retry/dead-letter behavior without failing the primary user request."),
            ("Failover integration test", "Verify that failover is rejected when the replication queue is not empty."),
            ("Failback integration test", "Verify maintenance mode, sync boundary call, and return to PRIMARY_ACTIVE."),
        ],
        [2.0, 5.6],
    )

    doc.add_section(WD_SECTION.CONTINUOUS)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
