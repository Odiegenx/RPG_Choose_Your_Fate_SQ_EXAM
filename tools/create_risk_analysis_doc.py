import os

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = "Risk_Analysis_SQL_Failover_Editable_Tables.docx"
TABLE_DIR = os.path.join("tools", "risk_analysis_tables")
TABLE_COUNTER = 0


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=8, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_width(cell, width_cm):
    width = Cm(width_cm)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width.twips)))


def set_cell_margins(cell, margin_twips=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn("w:" + side))
        if node is None:
            node = OxmlElement("w:" + side)
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2EC"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_grid(table, widths_cm):
    tbl = table._tbl
    existing_grid = tbl.tblGrid
    if existing_grid is not None:
        tbl.remove(existing_grid)

    tbl_grid = OxmlElement("w:tblGrid")
    for width_cm in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(Cm(width_cm).twips)))
        tbl_grid.append(grid_col)
    tbl.insert(1, tbl_grid)

    tbl_pr = tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def level_from_factor(factor):
    factor = int(factor)
    if factor >= 16:
        return "Very High"
    if factor >= 12:
        return "High"
    if factor >= 7:
        return "Medium"
    if factor >= 3:
        return "Low"
    return "Very Low"


def color_for_factor(factor):
    return RISK_COLOR_HEX[level_from_factor(factor)]


def table_cell_fill(headers, row_index, col_index, value):
    if row_index == 0:
        return "1F4E79"
    if headers[col_index] in ("Probability", "Impact"):
        return RISK_COLOR_HEX.get(str(value))
    if headers[col_index] == "Risk factor":
        return color_for_factor(value)
    if headers[0] == "Impact \\ Probability" and row_index > 0 and col_index > 0:
        first_value = str(value).split(":", 1)[0].strip()
        return color_for_factor(first_value)
    return "F7FAFC" if row_index % 2 == 0 else "FFFFFF"


def add_table(doc, headers, rows, widths=None, font_size=7):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    if not widths:
        widths = [1 for _ in headers]
    set_table_grid(table, widths)

    all_rows = [headers] + rows
    for row_index, values in enumerate(all_rows):
        if row_index > 0:
            row = table.add_row()
        else:
            row = table.rows[0]
            set_repeat_table_header(row)

        for col_index, value in enumerate(values):
            cell = row.cells[col_index]
            set_cell_width(cell, widths[col_index])
            set_cell_margins(cell, 110)
            shade_cell(cell, table_cell_fill(headers, row_index, col_index, value))
            is_header = row_index == 0
            text_color = (255, 255, 255) if is_header else (28, 35, 45)
            align = WD_ALIGN_PARAGRAPH.CENTER if headers[col_index] in ("ID", "P", "I", "Risk factor", "Found", "Review", "Status", "Owner") or is_header else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cell, value, bold=is_header, color=text_color, size=font_size, align=align)

    doc.add_paragraph()
    return table


def get_font(size, bold=False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        if os.path.exists(candidate):
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def wrap_text(draw, text, font, max_width):
    words = str(text).split()
    if not words:
        return [""]
    expanded_words = []
    for word in words:
        if draw.textbbox((0, 0), word, font=font)[2] <= max_width:
            expanded_words.append(word)
            continue
        chunk = ""
        for char in word:
            test = chunk + char
            if chunk and draw.textbbox((0, 0), test, font=font)[2] > max_width:
                expanded_words.append(chunk)
                chunk = char
            else:
                chunk = test
        if chunk:
            expanded_words.append(chunk)
    words = expanded_words
    lines = []
    current = words[0]
    for word in words[1:]:
        test = current + " " + word
        if draw.textbbox((0, 0), test, font=font)[2] <= max_width:
            current = test
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def render_table_image(headers, rows, widths_cm, image_path, font_size=8):
    scale = 38
    padding_x = 12
    padding_y = 10
    border = (200, 213, 225)
    header_fill = (31, 78, 121)
    stripe_fill = (247, 250, 252)
    white = (255, 255, 255)
    text = (28, 35, 45)

    if not widths_cm:
        widths_cm = [1 for _ in headers]
    col_widths = [max(55, int(w * scale)) for w in widths_cm]
    total_width = sum(col_widths)

    normal_font = get_font(max(11, int(font_size * 1.8)), False)
    header_font = get_font(max(11, int(font_size * 1.8)), True)
    line_gap = 4

    scratch = Image.new("RGB", (total_width, 100), white)
    draw = ImageDraw.Draw(scratch)

    all_rows = [headers] + rows
    row_heights = []
    wrapped_rows = []
    for row_index, row in enumerate(all_rows):
        font = header_font if row_index == 0 else normal_font
        wrapped_cells = []
        max_lines = 1
        for col_index, value in enumerate(row):
            lines = wrap_text(draw, value, font, col_widths[col_index] - (padding_x * 2))
            wrapped_cells.append(lines)
            max_lines = max(max_lines, len(lines))
        line_height = draw.textbbox((0, 0), "Ag", font=font)[3] + line_gap
        row_heights.append(max(42, padding_y * 2 + max_lines * line_height))
        wrapped_rows.append(wrapped_cells)

    total_height = sum(row_heights)
    image = Image.new("RGB", (total_width + 2, total_height + 2), white)
    draw = ImageDraw.Draw(image)

    y = 1
    for row_index, wrapped_cells in enumerate(wrapped_rows):
        x = 1
        row_height = row_heights[row_index]
        fill = header_fill if row_index == 0 else (stripe_fill if row_index % 2 == 0 else white)
        font = header_font if row_index == 0 else normal_font
        color = white if row_index == 0 else text
        for col_index, lines in enumerate(wrapped_cells):
            w = col_widths[col_index]
            cell_fill = fill
            raw_value = str(all_rows[row_index][col_index])
            header_value = str(headers[col_index])
            risk_key = raw_value.split("\n", 1)[0].split(":", 1)[0].strip()
            if row_index > 0 and (header_value in ("Level", "Risk level") or risk_key in RISK_COLORS):
                cell_fill = RISK_COLORS.get(risk_key, cell_fill)
            draw.rectangle([x, y, x + w, y + row_height], fill=cell_fill, outline=border, width=2)
            line_height = draw.textbbox((0, 0), "Ag", font=font)[3] + line_gap
            text_y = y + padding_y
            for line in lines:
                draw.text((x + padding_x, text_y), line, font=font, fill=color)
                text_y += line_height
            x += w
        y += row_height

    image.save(image_path)


def add_callout(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    title_run = p.add_run(title + ": ")
    title_run.bold = True
    title_run.font.size = Pt(10)
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    text_run = p.add_run(text)
    text_run.font.size = Pt(10)
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(31, 78, 121)
    return p


SCALE = {
    "Very Low": 1,
    "Low": 2,
    "Medium": 3,
    "High": 4,
    "Very High": 5,
}

RISK_COLORS = {
    "Very Low": (198, 239, 206),
    "Low": (226, 239, 218),
    "Medium": (255, 235, 156),
    "High": (244, 176, 132),
    "Very High": (248, 105, 107),
}

RISK_COLOR_HEX = {
    "Very Low": "C6EFCE",
    "Low": "E2EFDA",
    "Medium": "FFEB9C",
    "High": "F4B084",
    "Very High": "F8696B",
}


def risk_level(likelihood, impact):
    return level_from_factor(risk_factor(likelihood, impact))


def risk_factor(likelihood, impact):
    return scale_value(likelihood) * scale_value(impact)


def risk_matrix_value(probability, impact):
    return risk_factor(probability, impact)


def label_with_score(label):
    return f"{label} ({SCALE[label]})"


def scale_value(value):
    if isinstance(value, int):
        return value
    if isinstance(value, str) and value.isdigit():
        return int(value)
    return SCALE[value]


def score_label(score):
    return {
        1: "Very Low",
        2: "Low",
        3: "Medium",
        4: "High",
        5: "Very High",
    }[score]


def add_landscape_page_break(doc):
    doc.add_page_break()
    return "Low"


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.top_margin = Cm(1.4)
section.bottom_margin = Cm(1.4)
section.left_margin = Cm(1.4)
section.right_margin = Cm(1.4)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10)
for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
    styles[style_name].font.name = "Aptos Display"
    styles[style_name].font.color.rgb = RGBColor(31, 78, 121)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Risk Analysis")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(31, 78, 121)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("RPG Choose Your Fate - SQL Primary and Backup Database Failover")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(80, 80, 80)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("Software Quality Analysis | Focus: availability, reliability, recoverability, and project continuity")
run.italic = True
run.font.size = Pt(10)

doc.add_paragraph()
add_callout(
    doc,
    "Scope",
    "This analysis excludes MongoDB and Neo4j. The assessed architecture uses a primary SQL database and a secondary SQL backup database intended to take over automatically if the primary database becomes unavailable.",
)

add_heading(doc, "1. Introduction", 1)
doc.add_paragraph(
    "The system uses a primary SQL database as the main operational data source and a secondary SQL database as a backup/failover target. "
    "The backup database contains a copy of the primary database and is intended to keep the application available if the primary database fails. "
    "This improves availability, but it also introduces risks related to synchronization, failover detection, write consistency, recovery, and operational testing."
)
doc.add_paragraph(
    "The goal of this risk analysis is to identify the most important technical and general project risks, assess their probability and impact, and define mitigation and contingency plans. "
    "The analysis uses a five-level qualitative scale, supported by percentage ranges for probability, because it is detailed enough for software quality work while still being understandable in a project report."
)

add_heading(doc, "2. Risk Method", 1)
probability_rows = [
    ["1", "Very Low", "1-10%", "Event is highly unlikely to occur."],
    ["2", "Low", "11-30%", "Event is unlikely but possible."],
    ["3", "Medium", "31-60%", "Event may occur under certain conditions."],
    ["4", "High", "61-90%", "Event is likely to occur."],
    ["5", "Very High", "91-100%", "Event is almost certain to occur."],
]
impact_rows = [
    ["1", "Very Low", "Negligible impact on cost, schedule, and quality."],
    ["2", "Low", "Small increase in cost, minor delays, or small impact on quality."],
    ["3", "Medium", "Moderate cost increase or delays affecting milestones."],
    ["4", "High", "Major cost increase or delays impacting delivery."],
    ["5", "Very High", "Project cannot be completed or loses its business value."],
]
add_heading(doc, "2.1 Probability Scale", 2)
add_table(doc, ["Score", "Probability", "Range", "Meaning"], probability_rows, [2.0, 3.0, 3.0, 17.0], 8)
add_heading(doc, "2.2 Impact Scale", 2)
add_table(doc, ["Score", "Impact", "Meaning"], impact_rows, [2.0, 3.0, 19.0], 8)

add_heading(doc, "3. Risk Matrix", 1)
matrix_probabilities = [1, 2, 3, 4, 5]
matrix_impacts = [5, 4, 3, 2, 1]
identified_risks_for_matrix = [
    ("T1", 3, 4),
    ("T2", 3, 4),
    ("T3", 3, 4),
    ("T4", 3, 4),
    ("T5", 3, 4),
    ("T6", 3, 4),
    ("T7", 2, 4),
    ("T8", 2, 4),
    ("T9", 3, 3),
    ("T10", 3, 3),
    ("T11", 4, 4),
    ("T12", 3, 3),
    ("T13", 2, 4),
    ("T14", 3, 4),
    ("G1", 3, 3),
    ("G2", 2, 3),
    ("G3", 2, 4),
    ("G4", 3, 3),
    ("G5", 3, 4),
    ("G6", 3, 3),
    ("G7", 4, 3),
    ("G8", 3, 3),
    ("G9", 3, 3),
    ("G10", 3, 4),
]
matrix_lookup = {}
for risk_id, probability, impact in identified_risks_for_matrix:
    matrix_lookup.setdefault((probability, impact), []).append(risk_id)

def matrix_cell(probability, impact):
    factor = risk_matrix_value(probability, impact)
    ids = matrix_lookup.get((probability, impact), [])
    if ids:
        return f"{factor}: {', '.join(ids)}"
    return str(factor)

matrix_rows = [
    [str(impact)] + [matrix_cell(probability, impact) for probability in matrix_probabilities]
    for impact in matrix_impacts
]
add_table(
    doc,
    ["Impact \\ Probability"] + [str(value) for value in matrix_probabilities],
    matrix_rows,
    [4.2, 3.8, 3.8, 3.8, 3.8, 3.8],
    8,
)
doc.add_paragraph(
    "Risk factor is calculated as probability multiplied by impact. 1 = Very Low, 2 = Low, 3 = Medium, 4 = High, and 5 = Very High. Green cells indicate low exposure, yellow/orange cells indicate risks that should be actively monitored, and red cells indicate risks that require mitigation and a clear contingency plan."
)
doc.add_paragraph(
    "The scale uses scores from 1 to 5 for both probability and impact. The numeric risk factor is interpreted as follows: 1-2 = Very Low, 3-6 = Low, 7-11 = Medium, 12-15 = High, and 16-25 = Very High."
)

add_heading(doc, "4. Technical Risk Register", 1)
technical_headers = [
    "ID", "Area", "Risk, cause, and consequence", "P", "I", "Risk factor",
    "Found", "Review", "Status", "Mitigation, residual risk, and contingency", "Owner"
]
IDENTIFIED_DATE = "28/4/26"
FOLLOW_UP_DATE = "5/5/26"
technical_rows_raw = [
    ["T1", "Failover detection", "Risk: backup SQL is selected too late. Cause: health checks or retry logic are too slow. Consequence: users see downtime before failover.", 3, 4, "Mitigation: active DB health checks, short timeouts, clear thresholds. Contingency: manually route traffic to backup and restart backend.", "Backend", "Failover logic"],
    ["T2", "False failover", "Risk: backup is selected while primary is still usable. Cause: latency or aggressive timeout. Consequence: split-brain and inconsistent writes.", 3, 4, "Mitigation: conservative checks, retry windows, logged source switches. Contingency: force one active datasource and reconcile writes.", "BE/DevOps", "Health checks, logs"],
    ["T3", "Backup freshness", "Risk: backup SQL is outdated during failover. Cause: delayed or failed sync. Consequence: stale data or lost recent changes.", 3, 4, "Mitigation: monitor lag, row counts, checksums. Contingency: read-only mode or restore latest verified backup.", "Database", "Backup sync job"],
    ["T4", "Write divergence", "Risk: failover writes are not copied back to primary. Cause: backup is active while primary is offline. Consequence: DBs diverge.", 3, 4, "Mitigation: controlled failback and backup as temporary primary. Contingency: keep backup active until synchronization is verified.", "DB/BE", "Failback procedure"],
    ["T5", "Transaction consistency", "Risk: action succeeds in one DB but is missing in the other. Cause: async replication or failure during transaction. Consequence: inconsistent game/account state.", 3, 4, "Mitigation: transactional writes on active SQL and replication validation. Contingency: audit and replay missing writes if possible.", "BE/DB", "SQL transactions"],
    ["T6", "Automatic recovery", "Risk: system switches back to primary too early. Cause: primary returns stale/incomplete. Consequence: data loss or overwritten backup data.", 3, 4, "Mitigation: no automatic failback without sync validation. Contingency: keep backup active and rebuild primary.", "DevOps/DB", "Failback controls"],
    ["T7", "Authentication data", "Risk: credentials or roles differ. Cause: account changes are not synchronized. Consequence: login failure or wrong permissions.", 2, 4, "Mitigation: sync account/security tables with high priority. Contingency: validate against backup and review account changes.", "Security", "Account and role tables"],
    ["T8", "Authorization", "Risk: users access another user's data after failover. Cause: ownership data differs. Consequence: privacy/security breach.", 2, 4, "Mitigation: run ownership tests against both DBs. Contingency: disable affected endpoints until verified.", "Backend", "JWT, ownership services"],
    ["T9", "Configuration", "Risk: app points to wrong DB environment. Cause: wrong connection string/profile. Consequence: tests or demo use stale data.", 3, 3, "Mitigation: named profiles, clear env vars, startup datasource logging. Contingency: correct config and rerun smoke tests.", "DevOps", "application properties, Docker"],
    ["T10", "Performance", "Risk: failover logic slows requests. Cause: expensive health checks or retries. Consequence: poor UX and DB load.", 3, 3, "Mitigation: lightweight probes and cached health state. Contingency: temporary manual datasource override.", "Backend", "Health check implementation"],
    ["T11", "Testing", "Risk: failover is assumed but not proven. Cause: no realistic primary-failure test. Consequence: failure discovered during demo/use.", 4, 4, "Mitigation: stop primary, verify backup reads/writes, recover primary. Contingency: documented manual checklist.", "QA/Team", "Endpoint checklist, Docker"],
    ["T12", "Observability", "Risk: active SQL database is unclear. Cause: missing logs/status endpoint. Consequence: difficult debugging and weak evidence.", 3, 3, "Mitigation: log datasource switches and expose status. Contingency: inspect container logs and DB timestamps.", "Backend", "Logging, health endpoint"],
    ["T13", "Backup security", "Risk: backup SQL has weaker protection. Cause: secondary system gets weaker credentials. Consequence: sensitive data exposure.", 2, 4, "Mitigation: same credential and network policy as primary. Contingency: rotate credentials and audit DB access.", "Security", "Docker secrets/env"],
    ["T14", "Schema drift", "Risk: primary and backup schemas differ. Cause: migrations/scripts applied inconsistently. Consequence: runtime errors after failover.", 3, 4, "Mitigation: apply migrations to both DBs and compare schema. Contingency: rebuild backup and rerun scripts.", "Database", "SQL schema scripts"],
]
technical_rows = [
    row[:5] + [risk_factor(row[3], row[4]), IDENTIFIED_DATE, FOLLOW_UP_DATE, "Open", row[5], row[6]]
    for row in technical_rows_raw
]
technical_assessment_headers = ["ID", "Area", "P", "I", "Risk factor", "Found", "Review", "Status", "Owner"]
technical_detail_headers = ["ID", "Risk, cause, and consequence", "Mitigation, residual risk, and contingency"]
for start in range(0, len(technical_rows), 5):
    chunk = technical_rows[start:start + 5]
    assessment_rows = [[row[0], row[1], row[3], row[4], row[5], row[6], row[7], row[8], row[10]] for row in chunk]
    detail_rows = [[row[0], row[2], row[9]] for row in chunk]
    add_table(doc, technical_assessment_headers, assessment_rows, [1.0, 4.5, 1.0, 1.0, 2.5, 2.0, 2.0, 2.0, 4.2], 7)
    add_table(doc, technical_detail_headers, detail_rows, [1.0, 11.5, 12.0], 7)

add_heading(doc, "5. General Project Risk Register", 1)
general_headers = [
    "ID", "Area", "Risk, cause, and consequence", "P", "I", "Risk factor",
    "Found", "Review", "Status", "Mitigation and contingency", "Owner"
]
general_rows_raw = [
    ["G1", "Illness", "Risk: team member gets sick near deadline. Cause: illness/stress. Consequence: delayed tasks or unavailable knowledge.", 3, 3, "Mitigation: share knowledge and avoid single-person ownership. Contingency: reassign tasks and reduce scope.", "Team"],
    ["G2", "Power outage", "Risk: power failure blocks work/testing/demo prep. Cause: outage or uncharged laptop. Consequence: lost time.", 2, 3, "Mitigation: commit frequently and keep laptops charged. Contingency: move location or use another machine.", "Team"],
    ["G3", "Hardware failure", "Risk: laptop/dev machine fails. Cause: disk, charger, OS, or Docker problem. Consequence: cannot build/test.", 2, 4, "Mitigation: push code and document setup. Contingency: clone repository on another machine.", "Team"],
    ["G4", "Merge conflicts", "Risk: parallel work conflicts. Cause: same files or stale branches. Consequence: lost time and possible bugs.", 3, 3, "Mitigation: small branches and clear ownership. Contingency: resolve together and run smoke tests.", "Team"],
    ["G5", "Scope creep", "Risk: too many extra features. Cause: ambition or late changes. Consequence: core quality work suffers.", 3, 4, "Mitigation: prioritize and freeze optional features. Contingency: cut extras and focus on failover.", "Team/PM"],
    ["G6", "Roles", "Risk: tasks are assumed handled by someone else. Cause: no owner. Consequence: missing tests/docs/implementation.", 3, 3, "Mitigation: assign owners for backend, DB, tests, report. Contingency: short status meeting and redistribute tasks.", "Team/PM"],
    ["G7", "Deadline pressure", "Risk: quality decreases near hand-in. Cause: underestimated work or late bugs. Consequence: rushed testing.", 4, 3, "Mitigation: final-week checklist and feature freeze. Contingency: deliver smaller working solution.", "Team"],
    ["G8", "Lost files", "Risk: report/diagrams lost or overwritten. Cause: local-only work. Consequence: rework and confusion.", 3, 3, "Mitigation: version control/cloud storage. Contingency: recover from repository history or backups.", "Team"],
    ["G9", "External dependency", "Risk: Docker/DB image/package fails. Cause: version/cache issue. Consequence: setup or demo cannot run.", 3, 3, "Mitigation: pin versions and test on more than one machine. Contingency: use known working cache/machine.", "DevOps"],
    ["G10", "Presentation risk", "Risk: team cannot explain failover design. Cause: weak link to quality concepts. Consequence: lower assessment.", 3, 4, "Mitigation: prepare explanation linked to availability, reliability, recoverability, testability. Contingency: use risk analysis as notes.", "Team"],
]
general_rows = [
    row[:5] + [risk_factor(row[3], row[4]), IDENTIFIED_DATE, FOLLOW_UP_DATE, "Open"] + row[5:]
    for row in general_rows_raw
]
general_assessment_headers = ["ID", "Area", "P", "I", "Risk factor", "Found", "Review", "Status", "Owner"]
general_detail_headers = ["ID", "Risk, cause, and consequence", "Mitigation and contingency"]
for start in range(0, len(general_rows), 5):
    chunk = general_rows[start:start + 5]
    assessment_rows = [[row[0], row[1], row[3], row[4], row[5], row[6], row[7], row[8], row[10]] for row in chunk]
    detail_rows = [[row[0], row[2], row[9]] for row in chunk]
    add_table(doc, general_assessment_headers, assessment_rows, [1.0, 4.5, 1.0, 1.0, 2.5, 2.0, 2.0, 2.0, 4.2], 7)
    add_table(doc, general_detail_headers, detail_rows, [1.0, 11.5, 12.0], 7)

add_heading(doc, "6. Most Critical Risks", 1)
critical = [
    ("1. Backup SQL is outdated during failover", "This is the highest data-integrity risk. A backup database only improves availability if it is sufficiently synchronized with the primary database. If the backup is stale, users may see old game state, missing account changes, or outdated ownership information."),
    ("2. Write divergence after failover", "If the backup database accepts writes while the primary database is down, the system needs a safe way to merge or promote that data when the primary returns. Without a controlled failback process, the team risks overwriting newer data with older primary data."),
    ("3. False failover and split-brain behavior", "A temporary network issue can look like a database outage. If the system switches too aggressively, both databases may appear active in different parts of the system. This creates inconsistent behavior and makes debugging difficult."),
    ("4. Failover not tested realistically", "Failover is a quality feature only if it has been verified. The team should test by stopping the primary SQL database, confirming that the application switches to backup, and then testing recovery/failback behavior."),
    ("5. Security and authorization mismatch", "The backup contains the same sensitive data and must enforce the same account, role, and ownership rules. A failover system is not acceptable if it preserves availability but weakens access control."),
]
for title_text, body in critical:
    p = doc.add_paragraph()
    r = p.add_run(title_text)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)
    doc.add_paragraph(body)

doc.add_page_break()
add_heading(doc, "7. Recommended Verification Plan", 1)
verification_rows = [
    ["V1", "Normal operation", "Start application with primary SQL available.", "Application reads/writes from primary SQL."],
    ["V2", "Primary failure", "Stop primary SQL container/service while application is running.", "Application detects failure and switches to backup SQL."],
    ["V3", "Backup read", "Read characters/accounts after failover.", "Expected data is available from backup."],
    ["V4", "Backup write", "Create or update a test record during failover, if writes are allowed.", "Write succeeds on active backup or system enters controlled read-only mode."],
    ["V5", "Primary recovery", "Start primary SQL again.", "Application does not switch back until synchronization is verified."],
    ["V6", "Failback", "Synchronize data and switch back according to procedure.", "Primary and backup contain consistent data after recovery."],
    ["V7", "Security regression", "Repeat ownership/role tests against active backup.", "Users can still only access authorized data."],
    ["V8", "Observability", "Inspect logs/status endpoint during each step.", "Datasource changes are visible and explainable."],
]
add_table(doc, ["ID", "Test area", "Action", "Expected result"], verification_rows, [1.0, 4.0, 8.0, 12.0], 8)

doc.add_page_break()
add_heading(doc, "8. Residual Risk", 1)
doc.add_paragraph(
    "Even after mitigation, some residual risk remains. Automatic failover cannot guarantee zero downtime because the application must first detect the failure. "
    "A backup database cannot guarantee zero data loss unless synchronization is fully synchronous, which may reduce performance and increase system complexity. "
    "The safest design is therefore to combine automatic failover with clear logging, health checks, backup freshness validation, and a controlled failback process."
)
doc.add_paragraph(
    "From a software quality perspective, the backup SQL database improves availability and recoverability, but it also increases complexity. "
    "The team should present this as a trade-off: the design reduces the risk of complete service outage, while introducing new risks around consistency, monitoring, and operational discipline."
)

doc.add_page_break()
add_heading(doc, "9. Conclusion", 1)
doc.add_paragraph(
    "The most important risks in this project are not only implementation bugs, but risks related to how the system behaves under failure. "
    "The primary/backup SQL design supports higher availability, but it must be supported by synchronization checks, realistic failover tests, safe recovery procedures, and the same security controls on both databases. "
    "General project risks such as illness, power failure, hardware issues, and deadline pressure should also be managed because they can directly affect the team's ability to implement, test, and explain the quality solution."
)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("Risk Analysis - RPG Choose Your Fate")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(120, 120, 120)

doc.save(OUTPUT)
print(OUTPUT)
