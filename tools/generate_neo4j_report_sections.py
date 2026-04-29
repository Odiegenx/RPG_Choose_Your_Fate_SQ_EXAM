from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "Neo4j-report-sections-completed.docx"
OUT_MD = ROOT / "neo4j-report-sections-completed.md"
OUT_PNG = ROOT / "database-diagrams" / "neo4j_graph_model.png"


REPORT = [
    ("title", "Neo4j Report Sections"),
    ("p", "Draft-ready sections for chapter 4, the migration part of chapter 5, and the discussion in chapter 6."),
    ("h1", "4. Graph database"),
    ("p", "This chapter describes the graph database part of the project. The graph database is implemented with Neo4j and is used as an alternative representation of the same RPG domain that is stored in the relational MySQL database. The purpose is not only to copy data into another database engine, but to model the RPG domain in a form where relationships are first-class parts of the data model."),
    ("h2", "4.1. Intro to graph databases"),
    ("p", "A graph database stores data as nodes and relationships instead of tables and rows. Nodes represent domain entities, such as Account, Character, Scene, Choice, Quest, Item, NPC, RaceDetails and Inventory. Relationships represent how these entities are connected, for example an account owning a character, a scene having choices, or a choice leading to another scene."),
    ("p", "This model fits a choose-your-own-adventure RPG because the important questions are often relationship-oriented. The application needs to know which choices are available in a scene, where each choice leads, which quests and NPCs belong to a scene, which items are affected by choices, and which path a character has taken through the story. In a relational database these questions are answered through joins across multiple tables. In Neo4j the same questions can often be expressed as graph traversals, where the query follows stored relationships directly."),
    ("p", "The graph database is therefore especially useful for story flow, character progression and relationship-heavy queries. MySQL is still useful for strict relational consistency and stored database logic, but Neo4j gives a clearer representation of how the RPG world is connected."),
    ("h2", "4.2. Database design"),
    ("p", "The Neo4j model was designed by transforming the relational MySQL schema into a graph structure. Entity tables were mapped to node labels, while foreign keys and join tables were mapped to relationships. Some one-to-one tables were flattened into node properties, and the equipment table was represented as typed relationships to Item nodes where that better matched the graph model. The original MySQL ids are kept as node properties. This makes the migration deterministic and makes it easier to compare data across MySQL and Neo4j."),
    ("image", str(OUT_PNG)),
    ("caption", "Figure 4.1: Neo4j graph database model. This is a structure/model diagram, not a screenshot of stored data."),
    ("p", "Model overview:"),
    ("code", """(Account)-[:OWNS_CHARACTER]->(Character)
(Character)-[:CURRENT_SCENE]->(Scene)
(Character)-[:IN_CHAPTER]->(Chapter)
(Character)-[:HAS_RACE]->(RaceDetails)
(Character)-[:HAS_PATH]->(CharacterPath)
(Character)-[:HAS_INVENTORY]->(Inventory)
(Chapter)-[:HAS_SCENE]->(Scene)
(Scene)-[:HAS_CHOICE]->(Choice)-[:LEADS_TO]->(Scene)
(Scene)-[:HAS_QUEST]->(Quest)
(Scene)-[:HAS_NPC]->(Npc)-[:HAS_RACE]->(RaceDetails)
(Character)-[:HAS_QUEST {status}]->(Quest)
(CharacterPath)-[:INCLUDES_CHOICE]->(Choice)
(Inventory)-[:CONTAINS {amount}]->(Item)
(Quest)-[:INVOLVES_ITEM]->(Item)
(Choice)-[:AFFECTS_ITEM]->(Item)
(Character)-[:EQUIPPED_HEAD|EQUIPPED_CHEST|EQUIPPED_LEGS]->(Item)"""),
    ("table", [
        ["SQL source", "Neo4j representation", "Reason"],
        ["account", "(:Account)", "A user account is a domain entity and owns characters."],
        ["character_avatar", "(:Character)", "The player character connects to account, scene, chapter, race, inventory, quests and path."],
        ["chapter, scene", "(:Chapter), (:Scene), [:HAS_SCENE]", "Story locations are naturally connected as graph nodes."],
        ["choice", "(:Choice), [:HAS_CHOICE], [:LEADS_TO]", "Choices represent branching story traversal."],
        ["quest", "(:Quest), [:HAS_QUEST]", "Quests belong to scenes and can also be linked to characters and items."],
        ["item", "(:Item)", "Items are shared between inventory, quests, choices and equipment."],
        ["npc", "(:Npc), [:HAS_NPC], [:HAS_RACE]", "NPCs appear in scenes and can be related to race details."],
        ["inventory, inventory_has_item", "(:Inventory), [:CONTAINS {amount}]", "The join table becomes a relationship with an amount property."],
        ["character_path, character_path_choice", "(:CharacterPath), [:INCLUDES_CHOICE]", "The character history becomes a traversable path."],
        ["equipment", "[:EQUIPPED_HEAD], [:EQUIPPED_CHEST], [:EQUIPPED_LEGS]", "Equipment slots are clearer as typed relationships to Item nodes."],
        ["character_details", "Properties on (:Character)", "The one-to-one details table is flattened into the Character node."],
    ]),
    ("note", "Note: The table does not mean that every SQL table becomes a Neo4j node. Entity tables become nodes, while join tables and foreign-key dependencies are represented as relationships. Some one-to-one data is flattened into node properties when that gives a simpler graph model."),
    ("p", "The graph model separates independent nodes from dependent nodes. Independent nodes such as Account, Chapter, Item, RaceDetails and Npc can be created before every relationship is attached. Dependent nodes such as Character, CharacterPath and Inventory only make sense when they are connected to the surrounding domain context. A Character should be linked to an Account, a current Scene, a Chapter and RaceDetails. A CharacterPath and Inventory should be linked to a Character. This dependency model explains why the migration first creates nodes and then creates relationships."),
    ("todo", "TODO 4.2: Add final confirmation that the diagram matches the implemented Neo4j model."),
    ("h2", "4.3. Indexes, transactions, keys, constraints and stored objects"),
    ("p", "Neo4j does not use primary keys in the same way as a relational database. In this project, the original MySQL ids are stored as the id property on Neo4j nodes, and uniqueness constraints are created for the important node labels. These constraints act as stable identifiers for the graph model and support efficient lookup by id."),
    ("bullets", [
        "Unique constraints are created for Account, Chapter, Scene, Choice, Quest, Item, Npc, RaceDetails, Character, CharacterPath and Inventory.",
        "The constraints are created with CREATE CONSTRAINT ... REQUIRE n.id IS UNIQUE, which also gives Neo4j an index-backed way to find nodes by id.",
        "The migration uses MERGE instead of plain CREATE for the central nodes, so rerunning migration statements can update existing graph elements instead of always duplicating them.",
        "The Neo4j migration uses the Neo4j Java Driver. Node and relationship creation is executed inside a write transaction with session.executeWrite(...). If a write fails, the transaction can roll back instead of leaving a partly migrated graph.",
        "Integrity checks are executed in read transactions. These checks validate graph-specific assumptions such as characters having accounts, inventories belonging to characters, and scenes belonging to chapters.",
    ]),
    ("p", "The project does not define custom stored procedures or triggers inside Neo4j. In MySQL, stored procedures and triggers are used for lifecycle logic such as creating or deleting dependent character data. In the graph implementation those database-side mechanisms are replaced by application-layer services and explicit Cypher statements. This is a deliberate compromise: Neo4j supports transactions and constraints, but it does not enforce relational foreign keys or trigger-based lifecycle logic in the same way as MySQL."),
    ("todo", "TODO 4.3: Add actual constraint/index examples from Neo4j if needed."),
    ("todo", "TODO 4.3: Add a short note about where transactions are used in the code."),
    ("h2", "4.4. Description of the CRUD application for the graph database"),
    ("p", "The CRUD application uses the same controller-facing API style as the relational implementation. The backend can route requests to different data layers through the X-Data-Source header. If the header is missing, SQL is used as the default. If the header is set to neo4j or graph, the request is routed to the Neo4j data access implementation."),
    ("p", "The current Neo4j runtime CRUD implementation is most complete for Account. It uses an AccountNode model annotated with @Node(\"Account\"), an AccountNodeRepository extending Neo4jRepository, and a Neo4jAccountService that implements the same AccountDataAccess interface as the SQL account service. This lets AccountController expose one API contract while the actual data layer changes behind the service boundary."),
    ("p", "Account CRUD supports reading all accounts, reading by id, creating accounts, updating accounts and deleting accounts in Neo4j. The service also checks username and email uniqueness in the application layer before saving, encodes passwords, and maps AccountNode objects into AccountResponseDTO objects."),
    ("p", "Other Neo4j service classes exist for Character, CharacterDetails, CharacterPath, Equipment and RaceDetails, but these currently throw UnsupportedOperationException. They show the intended architecture, but they are not complete CRUD implementations yet. Therefore, the report should describe the Neo4j CRUD application as partially implemented: Account CRUD is functional, while the rest of the graph runtime layer is planned and would need equivalent Cypher/Spring Data Neo4j logic before it can fully match the SQL application."),
    ("todo", "TODO 4.4: Add final status of which Neo4j CRUD endpoints are implemented."),
    ("h1", "5. Migration"),
    ("p", "The Neo4j migration moves data from the relational MySQL database into a graph model. MySQL is treated as the source of truth, and the migration reads from the existing JPA repositories before transforming relational rows, foreign keys and join tables into Neo4j nodes and relationships."),
    ("p", "The Neo4j migration is exposed through POST /api/migrations/neo4j with a clearExisting request parameter. It also exposes GET /api/migrations/neo4j/integrity, which runs validation queries against the graph. The migration returns a response containing the target database name, whether existing graph data was cleared, the migration timestamp, migrated row counts and integrity violation counts."),
    ("p", "The Neo4j migration reads accounts, chapters, scenes, choices, quests, items, NPCs, race details, characters, character details, character paths, inventories and all relevant join-table data from MySQL using JPA repositories. It then writes the graph with explicit Cypher through the Neo4j Java Driver."),
    ("bullets", [
        "First, uniqueness constraints are created for central node labels.",
        "Second, if clearExisting is true, existing graph data is removed with MATCH (n) DETACH DELETE n.",
        "Third, nodes are created or updated with MERGE.",
        "Fourth, one-to-one character details are flattened onto Character nodes as properties.",
        "Fifth, relationships are created between existing nodes. Join tables become graph relationships, sometimes with relationship properties such as amount or status.",
        "Finally, graph integrity checks count missing required relationship patterns.",
    ]),
    ("p", "The Neo4j migration creates nodes before relationships. This avoids creating relationships to nodes that do not exist yet. For example, Account and Character nodes are created first, and then OWNS_CHARACTER relationships are created. The same pattern is used for chapters and scenes, scenes and choices, inventories and items, quests and items, and character paths and choices."),
    ("p", "The integrity checks compensate for the fact that Neo4j does not enforce relational foreign keys. The checks count graph problems such as characters without accounts, characters without scenes, characters without races, scenes without chapters, quests without scenes, inventories without characters and paths without characters. This makes data quality visible immediately after migration."),
    ("p", "A screenshot of the migrated graph can be included here as verification of the migration result. The Neo4j Browser result overview is useful because it shows the created node labels, such as Account, Character, Scene, Choice, Quest, Item, Npc, RaceDetails, Inventory and CharacterPath, together with relationship types such as OWNS_CHARACTER, CURRENT_SCENE, HAS_CHOICE, LEADS_TO and CONTAINS. This screenshot should be presented as evidence of the migrated data state, not as the database model. The database model itself is shown in Figure 4.1."),
    ("todo", "TODO 5: Insert Neo4j Browser screenshot of migrated graph data here."),
    ("todo", "TODO 5: Add migration result example with migrated counts and integrity check counts."),
    ("h1", "6. Discussion"),
    ("p", "This discussion focuses on the relational MySQL database compared with the Neo4j graph database. MySQL gives the strongest schema and integrity model, while Neo4j is strongest when the application needs to follow relationships through the story world."),
    ("h2", "Data modeling"),
    ("p", "MySQL models the domain as normalized tables with primary keys, foreign keys and join tables. This works well for structured data and avoids duplication, but relationship-heavy questions require joins. Neo4j models the same domain as nodes and relationships. This fits the story graph especially well because scenes, choices, quests, items, NPCs and character paths are naturally connected."),
    ("h2", "Transactions"),
    ("p", "MySQL uses transactions together with stored procedures and triggers for operations that must update several related tables. Neo4j also supports ACID transactions, and the graph migration uses write transactions for graph creation and read transactions for validation. The important difference is that Neo4j transactions make graph writes atomic, but they do not automatically enforce every relationship rule that a relational schema would enforce through foreign keys."),
    ("h2", "Queries"),
    ("p", "SQL is strong for precise filtering, reporting and joining normalized tables. Neo4j is strong for traversal queries, such as following choices from scene to scene, finding a character's path, or discovering connected quests, items and NPCs. For this RPG, the graph query model matches the narrative flow very naturally."),
    ("h2", "Performance"),
    ("p", "Performance depends on the access pattern. MySQL is efficient for indexed lookups and joins over a well-designed relational schema. Neo4j can be efficient for relationship traversal because relationships are stored directly and can be followed without repeatedly joining tables. Neo4j is not automatically faster for every query; it is strongest when the query shape follows relationships."),
    ("h2", "Constraints and integrity"),
    ("p", "MySQL provides the strongest built-in integrity through primary keys, foreign keys, constraints, stored procedures and triggers. Neo4j supports uniqueness constraints and indexes, but it does not enforce relationship cardinality like relational foreign keys. In this project, Neo4j integrity is handled through unique id constraints, ordered migration, application-layer logic and post-migration integrity checks."),
    ("h2", "Which database fits the domain best?"),
    ("p", "MySQL fits authentication, account data and strict consistency requirements well. Neo4j fits the story graph and progression model best because the core game experience is built from relationships: scenes have choices, choices lead to scenes, characters move through the story, and quests, NPCs and items are connected to that path."),
    ("h2", "Trade-offs and compromises"),
    ("p", "The main trade-off is that Neo4j gives a clear and flexible relationship model, but it requires more explicit integrity handling than MySQL. MySQL gives strong integrity, but the graph-like story structure is less visible because it is spread across tables and join tables. The project accepts this compromise by keeping a shared controller/service contract while allowing the Neo4j data layer to model relationship-heavy behavior in a way that fits a graph database."),
    ("todo", "TODO 6: Add one concrete Cypher query example to support the discussion."),
    ("todo", "TODO 6: Add performance measurements if any were made."),
    ("h2", "Sources"),
    ("bullets", [
        "Neo4j Cypher Manual, constraints: https://neo4j.com/docs/cypher-manual/current/schema/constraints/",
        "Neo4j Cypher Manual, creating constraints: https://neo4j.com/docs/cypher-manual/current/constraints/managing-constraints/",
        "Neo4j Operations Manual, transactions and ACID behavior: https://neo4j.com/docs/operations-manual/current/database-internals/transaction-management/",
    ]),
]


def load_font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def center(draw, xy, text, font, fill=(33, 37, 41)):
    x, y, w, h = xy
    lines = []
    for part in text.split("\n"):
        lines.extend(wrap(part, width=16) or [""])
    line_heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total_h = sum(line_heights) + (len(lines) - 1) * 5
    cy = y + (h - total_h) / 2
    for line, lh in zip(lines, line_heights):
        bbox = draw.textbbox((0, 0), line, font=font)
        tx = x + (w - (bbox[2] - bbox[0])) / 2
        draw.text((tx, cy), line, font=font, fill=fill)
        cy += lh + 5


def draw_arrow(draw, start, end, fill, width=4):
    draw.line([start, end], fill=fill, width=width)
    import math

    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 16
    p1 = (end[0] - size * math.cos(angle - 0.45), end[1] - size * math.sin(angle - 0.45))
    p2 = (end[0] - size * math.cos(angle + 0.45), end[1] - size * math.sin(angle + 0.45))
    draw.polygon([end, p1, p2], fill=fill)


def make_diagram():
    OUT_PNG.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1900, 1360), "#f8fafc")
    draw = ImageDraw.Draw(img)
    font = load_font(28)
    small = load_font(22)
    title = load_font(42, True)

    draw.text((60, 36), "Neo4j graph model - Choose Your Fate RPG", font=title, fill="#172033")
    draw.text((62, 92), "Labels and relationship types, not stored data", font=small, fill="#526071")

    heading = load_font(30, True)
    node_font = load_font(25, True)
    rel_font = load_font(24)

    def panel(x, y, w, h, label):
        draw.rounded_rectangle([x, y, x + w, y + h], radius=26, fill="#ffffff", outline="#cbd5e1", width=3)
        draw.text((x + 28, y + 22), label, font=heading, fill="#172033")

    def node_box(x, y, text, w=220):
        draw.rounded_rectangle([x, y, x + w, y + 72], radius=18, fill="#d9eafd", outline="#2b5c8a", width=3)
        center(draw, (x, y, w, 72), text, node_font)
        return (x, y, w, 72)

    def arrow_between(a, b, label, y_offset=0):
        ax, ay, aw, ah = a
        bx, by, bw, bh = b
        start = (ax + aw, ay + ah / 2 + y_offset)
        end = (bx, by + bh / 2 + y_offset)
        draw_arrow(draw, start, end, "#475569", 4)
        tb = draw.textbbox((0, 0), label, font=small)
        lx = (start[0] + end[0]) / 2 - (tb[2] - tb[0]) / 2
        ly = start[1] - 42
        draw.rounded_rectangle([lx - 8, ly - 6, lx + (tb[2] - tb[0]) + 8, ly + 28], radius=8, fill="#ffffff")
        draw.text((lx, ly), label, font=small, fill="#334155")

    panel(55, 150, 1790, 300, "Core story traversal")
    account = node_box(110, 275, "Account")
    character = node_box(410, 275, "Character", 240)
    scene = node_box(800, 275, "Scene")
    choice = node_box(1130, 275, "Choice")
    next_scene = node_box(1480, 275, "Scene")
    arrow_between(account, character, "OWNS_CHARACTER")
    arrow_between(character, scene, "CURRENT_SCENE")
    arrow_between(scene, choice, "HAS_CHOICE")
    arrow_between(choice, next_scene, "LEADS_TO")

    def relation_card(x, y, w, text, accent="#2b5c8a"):
        draw.rounded_rectangle([x, y, x + w, y + 58], radius=14, fill="#eef6ff", outline="#b6c9dd", width=2)
        draw.rectangle([x, y, x + 8, y + 58], fill=accent)
        draw.text((x + 22, y + 16), text, font=rel_font, fill="#26364a")

    panel(55, 490, 1790, 370, "Character aggregate")
    draw.text((105, 565), "Character node relationships", font=heading, fill="#172033")
    character_rels = [
        "Character -[:IN_CHAPTER]-> Chapter",
        "Character -[:HAS_RACE]-> RaceDetails",
        "Character -[:HAS_PATH]-> CharacterPath",
        "Character -[:HAS_INVENTORY]-> Inventory",
        "Character -[:HAS_QUEST {status}]-> Quest",
        "Character -[:EQUIPPED_HEAD|CHEST|LEGS]-> Item",
    ]
    for idx, rel in enumerate(character_rels):
        col = idx % 2
        row = idx // 2
        relation_card(105 + col * 820, 620 + row * 68, 760, rel)

    draw.text((105, 825), "Flattened one-to-one table: CharacterDetails -> Character properties (intelligence, charisma, fashion)", font=small, fill="#334155")

    panel(55, 900, 1790, 335, "World, story and item relationships")
    world_rels = [
        "Chapter -[:HAS_SCENE]-> Scene",
        "Scene -[:HAS_NPC]-> Npc",
        "Npc -[:HAS_RACE]-> RaceDetails",
        "Scene -[:HAS_QUEST]-> Quest",
        "Quest -[:INVOLVES_ITEM]-> Item",
        "Choice -[:AFFECTS_ITEM]-> Item",
        "Inventory -[:CONTAINS {amount}]-> Item",
        "CharacterPath -[:INCLUDES_CHOICE]-> Choice",
    ]
    for idx, rel in enumerate(world_rels):
        col = idx % 2
        row = idx // 2
        relation_card(105 + col * 820, 980 + row * 68, 760, rel, "#475569")

    img.save(OUT_PNG)


def write_markdown():
    lines = []
    for kind, value in REPORT:
        if kind == "title":
            lines.append(f"# {value}\n")
        elif kind == "h1":
            lines.append(f"\n# {value}\n")
        elif kind == "h2":
            lines.append(f"\n## {value}\n")
        elif kind == "p":
            lines.append(f"{value}\n")
        elif kind == "note":
            lines.append(f"> **{value}**\n")
        elif kind == "todo":
            lines.append(f"> **{value}**\n")
        elif kind == "image":
            rel = Path(value).relative_to(ROOT)
            lines.append(f"![Neo4j graph model]({rel.as_posix()})\n")
        elif kind == "caption":
            lines.append(f"*{value}*\n")
        elif kind == "code":
            lines.append(f"```text\n{value}\n```\n")
        elif kind == "bullets":
            lines.extend([f"- {item}" for item in value])
            lines.append("")
        elif kind == "table":
            header = value[0]
            lines.append("| " + " | ".join(header) + " |")
            lines.append("| " + " | ".join(["---"] * len(header)) + " |")
            for row in value[1:]:
                lines.append("| " + " | ".join(row) + " |")
            lines.append("")
    OUT_MD.write_text("\n".join(lines), encoding="utf-8")


def add_code_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def write_docx():
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    for kind, value in REPORT:
        if kind == "title":
            p = doc.add_heading(value, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "h1":
            doc.add_heading(value, level=1)
        elif kind == "h2":
            doc.add_heading(value, level=2)
        elif kind == "p":
            doc.add_paragraph(value)
        elif kind == "note":
            p = doc.add_paragraph()
            run = p.add_run(value)
            run.bold = True
            run.font.color.rgb = RGBColor(47, 84, 150)
            p.style = "Intense Quote"
        elif kind == "todo":
            p = doc.add_paragraph()
            run = p.add_run(value)
            run.bold = True
            run.font.color.rgb = RGBColor(192, 86, 33)
            p.style = "Intense Quote"
        elif kind == "image":
            doc.add_picture(value, width=Inches(6.7))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "caption":
            p = doc.add_paragraph(value)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].italic = True
            p.runs[0].font.size = Pt(9)
        elif kind == "code":
            add_code_paragraph(doc, value)
        elif kind == "bullets":
            for item in value:
                doc.add_paragraph(item, style="List Bullet")
        elif kind == "table":
            table = doc.add_table(rows=1, cols=len(value[0]))
            table.style = "Table Grid"
            for i, cell in enumerate(table.rows[0].cells):
                cell.text = value[0][i]
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for row in value[1:]:
                cells = table.add_row().cells
                for i, cell_text in enumerate(row):
                    cells[i].text = cell_text
    try:
        doc.save(OUT_DOCX)
    except PermissionError:
        fallback = OUT_DOCX.with_name(f"{OUT_DOCX.stem}-with-todos{OUT_DOCX.suffix}")
        doc.save(fallback)
        print(f"Could not overwrite locked file. Wrote {fallback}")


def main():
    make_diagram()
    write_markdown()
    write_docx()
    print(f"Wrote {OUT_DOCX}")
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_PNG}")


if __name__ == "__main__":
    main()
